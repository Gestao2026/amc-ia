#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Mapa do Edital. Gerador do checklist de edital em Word, no padrao visual da Mobilizando.

Le um markdown estruturado (com bloco de metadados no topo) e produz um .docx
com capa, sumario, cabecalho, rodape e as secoes formatadas.

Uso:
    python scripts/mapa-edital.py <caminho-do-markdown> [--saida <pasta ou arquivo .docx>]

O markdown aceita:
    ## N. TITULO DA SECAO      -> faixa verde numerada
    ### Subtitulo              -> subtitulo verde com fio dourado
    | tabela | markdown |      -> quadro com cabecalho verde
    - [ ] item                 -> item com quadradinho para marcar
    - item                     -> marcador comum
    1. item                    -> lista numerada
    > texto                    -> caixa de destaque (dourada)
    !> texto                   -> caixa de alerta (vermelha)
    ---                        -> fio dourado
    **negrito**, *italico*, `mono`

Bloco de metadados, nas primeiras linhas do arquivo, antes de qualquer secao:

    ---meta
    titulo: Nome curto do edital
    subtitulo: Linha de identificacao completa
    orgao: Orgao responsavel
    prazo: 31/08/2026 as 17h
    plataforma: Mapa Cultural BH
    data: 14/08/2026
    arquivo: mapa-edital-pnab-ciclo-2-bh-04-2026
    ---
"""

import argparse
import os
import re
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ---------------------------------------------------------------------------
# Identidade visual da Mobilizando Impacto Social
# ---------------------------------------------------------------------------

# Codigos oficiais, conferidos no Manual de Marca da Mobilizando (BrianLucz Studios)
# e nos pixels do proprio arquivo do logo.
VERDE_PROFUNDO = "1D624D"   # verde da marca. Capa, faixas de secao, cabecalho de quadro
VERDE = "1D624D"            # subtitulos e destaques de texto
VERDE_CLARO = "EAF1EE"      # tint do verde, preenchimento alternado de linha
DOURADO = "D1B484"          # dourado da marca. Fios, numeros de secao, detalhes
DOURADO_CLARO = "F7F0E4"    # tint do dourado, fundo das caixas de destaque
ROXO = "7F126F"             # terceira cor da paleta, usada nas caixas de alerta
ROXO_CLARO = "F7EDF5"
AZUL = "070552"             # quarta cor da paleta, reserva
CINZA_TEXTO = "2B2B2B"      # corpo de texto
CINZA_BORDA = "DCE4E1"
CINZA_APOIO = "6B7772"
BRANCO = "FFFFFF"

LARGURA_UTIL = Cm(17.0)   # 21 cm de papel menos as duas margens de 2 cm

# O Manual de Marca pede Poppins (titulos) e Inter (corpo). Nenhuma das duas esta
# instalada nesta maquina, e fonte ausente vira substituicao imprevisivel no Word do
# cliente. Ficamos em Calibri ate as fontes da marca serem instaladas nos dois lados.
FONTE = "Calibri"
FONTE_TITULO = "Calibri"

MARCA = "Mobilizando Impacto Social"
TAGLINE = "Impacto que Transforma"
NOME_DOCUMENTO = "MAPA DO EDITAL"

CAMINHOS_LOGO = [
    os.path.join("painel", "marca", "mobilizando-logo.png"),
    os.path.join("painel", "marca", "mobilizando-logo.jpg"),
    os.path.join("painel", "marca", "logo.png"),
]

# Versao vazada (branco e dourado), para assentar sobre o verde da capa
CAMINHOS_LOGO_ESCURO = [
    os.path.join("painel", "marca", "mobilizando-logo-fundo-escuro.png"),
]


# ---------------------------------------------------------------------------
# Utilitarios de baixo nivel do Word
# ---------------------------------------------------------------------------


def sombrear(elemento, cor_hex):
    """Aplica cor de preenchimento a uma celula ou paragrafo."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), cor_hex)
    elemento.append(shd)


def sombrear_celula(celula, cor_hex):
    sombrear(celula._tc.get_or_add_tcPr(), cor_hex)


def sombrear_paragrafo(paragrafo, cor_hex):
    sombrear(paragrafo._p.get_or_add_pPr(), cor_hex)


def borda(elemento_pr, lados, cor_hex, tamanho=8, tipo="single"):
    """Define bordas. `lados` e uma lista entre top, bottom, left, right."""
    tag = "w:pBdr" if elemento_pr.tag.endswith("pPr") else "w:tcBorders"
    existente = elemento_pr.find(qn(tag))
    if existente is None:
        existente = OxmlElement(tag)
        elemento_pr.append(existente)
    for lado in lados:
        el = existente.find(qn("w:" + lado))
        if el is None:
            el = OxmlElement("w:" + lado)
            existente.append(el)
        el.set(qn("w:val"), tipo)
        el.set(qn("w:sz"), str(tamanho))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), cor_hex)


def borda_paragrafo(paragrafo, lados, cor_hex, tamanho=8):
    borda(paragrafo._p.get_or_add_pPr(), lados, cor_hex, tamanho)


def borda_celula(celula, lados, cor_hex, tamanho=6):
    borda(celula._tc.get_or_add_tcPr(), lados, cor_hex, tamanho)


def margens_celula(celula, topo=80, base=80, esq=120, dir_=120):
    tcPr = celula._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for nome, valor in (("top", topo), ("bottom", base), ("left", esq), ("right", dir_)):
        el = OxmlElement("w:" + nome)
        el.set(qn("w:w"), str(valor))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def largura_tabela(tabela, largura):
    """Fixa a largura da tabela, para a faixa casar com os quadros."""
    tblPr = tabela._tbl.tblPr
    w = OxmlElement("w:tblW")
    w.set(qn("w:w"), str(int(largura.twips)))
    w.set(qn("w:type"), "dxa")
    tblPr.append(w)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)


def largura_total_pct(tabela):
    """Faz a tabela ocupar 100% da area util, mantendo as colunas automaticas."""
    w = OxmlElement("w:tblW")
    w.set(qn("w:w"), "5000")
    w.set(qn("w:type"), "pct")
    tabela._tbl.tblPr.append(w)


def sem_bordas(tabela):
    tblPr = tabela._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for lado in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement("w:" + lado)
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        borders.append(el)
    tblPr.append(borders)


def campo(paragrafo, instrucao):
    """Insere um campo do Word (usado para numeracao de pagina)."""
    run = paragrafo.add_run()
    ini = OxmlElement("w:fldChar")
    ini.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instrucao
    fim = OxmlElement("w:fldChar")
    fim.set(qn("w:fldCharType"), "end")
    run._r.append(ini)
    run._r.append(instr)
    run._r.append(fim)
    return run


def espacamento(paragrafo, antes=0, depois=6, entre_linhas=1.15):
    pf = paragrafo.paragraph_format
    pf.space_before = Pt(antes)
    pf.space_after = Pt(depois)
    pf.line_spacing = entre_linhas


# ---------------------------------------------------------------------------
# Texto com formatacao inline
# ---------------------------------------------------------------------------

PADRAO_INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")


def escrever_inline(paragrafo, texto, tamanho=10.5, cor=CINZA_TEXTO, negrito_base=False):
    """Escreve texto tratando **negrito**, *italico* e `mono`."""
    texto = re.sub(r"\[\[([^\]]+)\]\]", r"\1", texto)
    for parte in PADRAO_INLINE.split(texto):
        if not parte:
            continue
        run = paragrafo.add_run()
        if parte.startswith("**") and parte.endswith("**"):
            run.text = parte[2:-2]
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(VERDE_PROFUNDO)
        elif parte.startswith("*") and parte.endswith("*"):
            run.text = parte[1:-1]
            run.italic = True
            run.font.color.rgb = RGBColor.from_string(cor)
        elif parte.startswith("`") and parte.endswith("`"):
            run.text = parte[1:-1]
            run.font.name = "Consolas"
            run.font.color.rgb = RGBColor.from_string(VERDE)
        else:
            run.text = parte
            run.bold = negrito_base
            run.font.color.rgb = RGBColor.from_string(cor)
        if run.font.name is None:
            run.font.name = FONTE
        run.font.size = Pt(tamanho)


# ---------------------------------------------------------------------------
# Leitura do markdown
# ---------------------------------------------------------------------------


def ler_fonte(caminho):
    with open(caminho, "r", encoding="utf-8") as f:
        bruto = f.read()

    meta = {}
    corpo = bruto
    if bruto.lstrip().startswith("---meta"):
        inicio = bruto.index("---meta") + len("---meta")
        fim = bruto.index("\n---", inicio)
        for linha in bruto[inicio:fim].strip().splitlines():
            if ":" in linha:
                chave, valor = linha.split(":", 1)
                meta[chave.strip().lower()] = valor.strip()
        corpo = bruto[fim + len("\n---"):]
    return meta, corpo.strip()


def linha_tabela(linha):
    return linha.strip().startswith("|") and linha.strip().endswith("|")


def celulas_de(linha):
    return [c.strip() for c in linha.strip().strip("|").split("|")]


def eh_separador(linha):
    return bool(re.fullmatch(r"\|[\s:\-|]+\|", linha.strip()))


# ---------------------------------------------------------------------------
# Blocos do documento
# ---------------------------------------------------------------------------


def achar_logo(raiz, caminhos=None):
    for rel in (caminhos or CAMINHOS_LOGO):
        caminho = os.path.join(raiz, rel)
        if os.path.exists(caminho):
            return caminho
    return None


def faixa(doc, texto_esq, texto_dir="", cor=VERDE_PROFUNDO, altura_pt=11):
    """Faixa colorida de largura total, usada nos titulos de secao.

    Uma celula so, com o numero da secao empurrado para a direita por tabulacao.
    Duas celulas deixariam um fio vertical visivel na emenda, mesmo sem borda.
    """
    tabela = doc.add_table(rows=1, cols=1)
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabela.autofit = False
    sem_bordas(tabela)
    celula = tabela.rows[0].cells[0]
    celula.width = LARGURA_UTIL
    largura_tabela(tabela, LARGURA_UTIL)
    sombrear_celula(celula, cor)
    margens_celula(celula, topo=110, base=110, esq=200, dir_=200)

    p = celula.paragraphs[0]
    espacamento(p, 0, 0, 1.0)
    if texto_dir:
        p.paragraph_format.tab_stops.add_tab_stop(
            LARGURA_UTIL - Cm(0.7), WD_TAB_ALIGNMENT.RIGHT
        )
    run = p.add_run(texto_esq)
    run.bold = True
    run.font.size = Pt(altura_pt)
    run.font.name = FONTE_TITULO
    run.font.color.rgb = RGBColor.from_string(BRANCO)
    if texto_dir:
        r2 = p.add_run("\t" + texto_dir)
        r2.bold = True
        r2.font.size = Pt(altura_pt)
        r2.font.name = FONTE_TITULO
        r2.font.color.rgb = RGBColor.from_string(DOURADO)
    depois = doc.add_paragraph()
    espacamento(depois, 0, 2, 1.0)
    return tabela


def fio(doc, cor=DOURADO, tamanho=12, espaco_antes=4, espaco_depois=8):
    p = doc.add_paragraph()
    espacamento(p, espaco_antes, espaco_depois, 1.0)
    p.paragraph_format.space_after = Pt(espaco_depois)
    borda_paragrafo(p, ["bottom"], cor, tamanho)
    return p


def caixa(doc, linhas, cor_fundo, cor_borda, rotulo=None):
    tabela = doc.add_table(rows=1, cols=1)
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabela.autofit = False
    sem_bordas(tabela)
    largura_tabela(tabela, LARGURA_UTIL)
    celula = tabela.rows[0].cells[0]
    celula.width = LARGURA_UTIL
    sombrear_celula(celula, cor_fundo)
    margens_celula(celula, topo=140, base=140, esq=220, dir_=220)
    borda_celula(celula, ["left"], cor_borda, tamanho=24)
    celula.paragraphs[0]._p.getparent().remove(celula.paragraphs[0]._p)
    if rotulo:
        p = celula.add_paragraph()
        espacamento(p, 0, 3, 1.0)
        r = p.add_run(rotulo)
        r.bold = True
        r.font.size = Pt(9)
        r.font.name = FONTE
        r.font.color.rgb = RGBColor.from_string(cor_borda)
    for i, linha in enumerate(linhas):
        p = celula.add_paragraph()
        espacamento(p, 0, 2 if i < len(linhas) - 1 else 0, 1.15)
        escrever_inline(p, linha, tamanho=10)
    depois = doc.add_paragraph()
    espacamento(depois, 0, 4, 1.0)


def quadro(doc, linhas):
    """Renderiza uma tabela markdown como quadro formatado."""
    dados = [celulas_de(l) for l in linhas if not eh_separador(l)]
    if not dados:
        return
    n_col = max(len(l) for l in dados)
    dados = [l + [""] * (n_col - len(l)) for l in dados]

    tabela = doc.add_table(rows=len(dados), cols=n_col)
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabela.autofit = True
    sem_bordas(tabela)
    largura_total_pct(tabela)

    for i, linha in enumerate(dados):
        cabecalho = i == 0
        for j, valor in enumerate(linha):
            celula = tabela.rows[i].cells[j]
            margens_celula(celula, topo=70, base=70, esq=120, dir_=120)
            p = celula.paragraphs[0]
            espacamento(p, 0, 0, 1.05)
            if cabecalho:
                sombrear_celula(celula, VERDE_PROFUNDO)
                r = p.add_run(valor.replace("**", ""))
                r.bold = True
                r.font.size = Pt(9.5)
                r.font.name = FONTE
                r.font.color.rgb = RGBColor.from_string(BRANCO)
            else:
                if i % 2 == 0:
                    sombrear_celula(celula, VERDE_CLARO)
                borda_celula(celula, ["top"], CINZA_BORDA, tamanho=4)
                escrever_inline(p, valor, tamanho=9.5)
    depois = doc.add_paragraph()
    espacamento(depois, 2, 6, 1.0)


def item_marcavel(doc, texto):
    p = doc.add_paragraph()
    espacamento(p, 0, 3, 1.15)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    marca = p.add_run("\u2751  ")
    marca.font.size = Pt(11)
    marca.font.name = FONTE
    marca.font.color.rgb = RGBColor.from_string(DOURADO)
    escrever_inline(p, texto, tamanho=10.5)


def marcador(doc, texto, nivel=0):
    p = doc.add_paragraph()
    espacamento(p, 0, 3, 1.15)
    p.paragraph_format.left_indent = Cm(0.6 + 0.5 * nivel)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    ponto = p.add_run("\u25AA  " if nivel == 0 else "\u2013  ")
    ponto.font.size = Pt(9)
    ponto.font.name = FONTE
    ponto.font.color.rgb = RGBColor.from_string(VERDE)
    escrever_inline(p, texto, tamanho=10.5)


def numerado(doc, numero, texto):
    p = doc.add_paragraph()
    espacamento(p, 0, 3, 1.15)
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.8)
    n = p.add_run(f"{numero}.  ")
    n.bold = True
    n.font.size = Pt(10.5)
    n.font.name = FONTE
    n.font.color.rgb = RGBColor.from_string(DOURADO)
    escrever_inline(p, texto, tamanho=10.5)


def subtitulo(doc, texto):
    p = doc.add_paragraph()
    espacamento(p, 8, 4, 1.0)
    r = p.add_run(texto)
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = FONTE_TITULO
    r.font.color.rgb = RGBColor.from_string(VERDE)
    borda_paragrafo(p, ["bottom"], DOURADO, 6)


def sub_subtitulo(doc, texto):
    p = doc.add_paragraph()
    espacamento(p, 6, 2, 1.0)
    r = p.add_run(texto)
    r.bold = True
    r.font.size = Pt(10)
    r.font.name = FONTE_TITULO
    r.font.color.rgb = RGBColor.from_string(CINZA_TEXTO)


def paragrafo_comum(doc, texto):
    p = doc.add_paragraph()
    espacamento(p, 0, 5, 1.2)
    escrever_inline(p, texto, tamanho=10.5)


# ---------------------------------------------------------------------------
# Capa, sumario, cabecalho e rodape
# ---------------------------------------------------------------------------


def montar_capa(doc, meta, logo_escuro):
    bloco = doc.add_table(rows=1, cols=1)
    bloco.autofit = False
    sem_bordas(bloco)
    largura_tabela(bloco, LARGURA_UTIL)
    celula = bloco.rows[0].cells[0]
    celula.width = LARGURA_UTIL
    sombrear_celula(celula, VERDE_PROFUNDO)
    margens_celula(celula, topo=420, base=420, esq=340, dir_=340)
    celula.paragraphs[0]._p.getparent().remove(celula.paragraphs[0]._p)

    if logo_escuro:
        p = celula.add_paragraph()
        espacamento(p, 0, 14, 1.0)
        p.add_run().add_picture(logo_escuro, width=Cm(5.4))
    else:
        p = celula.add_paragraph()
        espacamento(p, 0, 2, 1.0)
        r = p.add_run(MARCA.upper())
        r.bold = True
        r.font.size = Pt(15)
        r.font.name = FONTE_TITULO
        r.font.color.rgb = RGBColor.from_string(BRANCO)
        p2 = celula.add_paragraph()
        espacamento(p2, 0, 14, 1.0)
        r2 = p2.add_run(TAGLINE)
        r2.font.size = Pt(9.5)
        r2.font.name = FONTE
        r2.font.color.rgb = RGBColor.from_string(DOURADO)

    p = celula.add_paragraph()
    espacamento(p, 0, 4, 1.0)
    r = p.add_run(NOME_DOCUMENTO)
    r.bold = True
    r.font.size = Pt(30)
    r.font.name = FONTE_TITULO
    r.font.color.rgb = RGBColor.from_string(BRANCO)

    p = celula.add_paragraph()
    espacamento(p, 0, 12, 1.0)
    r = p.add_run("Checklist completo para conferir antes de inscrever")
    r.font.size = Pt(10)
    r.font.name = FONTE
    r.font.color.rgb = RGBColor.from_string(DOURADO)

    linha_ouro = celula.add_paragraph()
    espacamento(linha_ouro, 0, 12, 1.0)
    borda_paragrafo(linha_ouro, ["bottom"], DOURADO, 16)

    p = celula.add_paragraph()
    espacamento(p, 0, 4, 1.15)
    r = p.add_run(meta.get("titulo", ""))
    r.bold = True
    r.font.size = Pt(16)
    r.font.name = FONTE_TITULO
    r.font.color.rgb = RGBColor.from_string(BRANCO)

    if meta.get("subtitulo"):
        p = celula.add_paragraph()
        espacamento(p, 0, 0, 1.2)
        r = p.add_run(meta["subtitulo"])
        r.font.size = Pt(10.5)
        r.font.name = FONTE
        r.font.color.rgb = RGBColor.from_string("D9E4DE")

    doc.add_paragraph()

    # Quadro de identificacao rapida logo abaixo da capa
    campos = [
        ("Órgão responsável", meta.get("orgao", "")),
        ("Prazo final de inscrição", meta.get("prazo", "")),
        ("Onde se inscreve", meta.get("plataforma", "")),
        ("Documento preparado em", meta.get("data", "")),
    ]
    campos = [(r, v) for r, v in campos if v]
    tabela = doc.add_table(rows=len(campos), cols=2)
    tabela.autofit = False
    sem_bordas(tabela)
    largura_tabela(tabela, LARGURA_UTIL)
    for i, (rotulo, valor) in enumerate(campos):
        c1, c2 = tabela.rows[i].cells
        c1.width = Cm(5.4)
        c2.width = Cm(11.6)
        for c in (c1, c2):
            margens_celula(c, topo=90, base=90, esq=140, dir_=140)
            sombrear_celula(c, VERDE_CLARO if i % 2 == 0 else BRANCO)
            borda_celula(c, ["top"], CINZA_BORDA, tamanho=4)
        p1 = c1.paragraphs[0]
        espacamento(p1, 0, 0, 1.0)
        r1 = p1.add_run(rotulo.upper())
        r1.bold = True
        r1.font.size = Pt(8.5)
        r1.font.name = FONTE
        r1.font.color.rgb = RGBColor.from_string(VERDE)
        p2 = c2.paragraphs[0]
        espacamento(p2, 0, 0, 1.0)
        escrever_inline(p2, valor, tamanho=10)

    rodape_capa = doc.add_paragraph()
    espacamento(rodape_capa, 16, 0, 1.2)
    r = rodape_capa.add_run(
        "Este documento reúne, em um lugar só, tudo que o edital exige. "
        "Cada quadradinho é um item para conferir. O que ficar sem marcar é o que ainda falta resolver."
    )
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.name = FONTE
    r.font.color.rgb = RGBColor.from_string("5A6660")


def montar_sumario(doc, secoes):
    faixa(doc, "SUMÁRIO")
    tabela = doc.add_table(rows=len(secoes), cols=2)
    tabela.autofit = False
    sem_bordas(tabela)
    largura_tabela(tabela, LARGURA_UTIL)
    for i, (numero, titulo) in enumerate(secoes):
        c1, c2 = tabela.rows[i].cells
        c1.width = Cm(1.6)
        c2.width = Cm(15.4)
        for c in (c1, c2):
            margens_celula(c, topo=60, base=60, esq=100, dir_=100)
            borda_celula(c, ["bottom"], CINZA_BORDA, tamanho=4)
        p1 = c1.paragraphs[0]
        espacamento(p1, 0, 0, 1.0)
        r1 = p1.add_run(numero)
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.name = FONTE
        r1.font.color.rgb = RGBColor.from_string(DOURADO)
        p2 = c2.paragraphs[0]
        espacamento(p2, 0, 0, 1.0)
        r2 = p2.add_run(titulo)
        r2.font.size = Pt(10)
        r2.font.name = FONTE
        r2.font.color.rgb = RGBColor.from_string(CINZA_TEXTO)


def montar_cabecalho_rodape(doc, meta, logo):
    secao = doc.sections[0]
    secao.different_first_page_header_footer = True

    cab = secao.header
    p = cab.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    espacamento(p, 0, 4, 1.0)
    if logo:
        p.paragraph_format.tab_stops.add_tab_stop(LARGURA_UTIL, WD_TAB_ALIGNMENT.RIGHT)
    r = p.add_run(NOME_DOCUMENTO + "   ")
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.name = FONTE
    r.font.color.rgb = RGBColor.from_string(VERDE_PROFUNDO)
    r2 = p.add_run(meta.get("titulo", ""))
    r2.font.size = Pt(8.5)
    r2.font.name = FONTE
    r2.font.color.rgb = RGBColor.from_string(CINZA_APOIO)
    if logo:
        p.add_run("	").add_picture(logo, width=Cm(2.2))
    borda_paragrafo(p, ["bottom"], DOURADO, 6)

    rod = secao.footer
    p = rod.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    espacamento(p, 4, 0, 1.0)
    borda_paragrafo(p, ["top"], CINZA_BORDA, 4)
    r = p.add_run(f"{MARCA}  |  {TAGLINE}  |  página ")
    r.font.size = Pt(8)
    r.font.name = FONTE
    r.font.color.rgb = RGBColor.from_string(CINZA_APOIO)
    rp = campo(p, "PAGE")
    rp.font.size = Pt(8)
    rp.font.name = FONTE
    rp.font.color.rgb = RGBColor.from_string(VERDE)

    rod_primeira = secao.first_page_footer
    p = rod_primeira.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    espacamento(p, 0, 0, 1.0)
    r = p.add_run(f"{MARCA}  |  {TAGLINE}")
    r.font.size = Pt(8)
    r.font.name = FONTE
    r.font.color.rgb = RGBColor.from_string(CINZA_APOIO)


# ---------------------------------------------------------------------------
# Renderizacao do corpo
# ---------------------------------------------------------------------------


def coletar_secoes(corpo):
    secoes = []
    for linha in corpo.splitlines():
        if linha.startswith("## "):
            titulo = linha[3:].strip()
            m = re.match(r"^(\d+)\.\s*(.+)$", titulo)
            if m:
                secoes.append((m.group(1), m.group(2)))
            else:
                secoes.append(("", titulo))
    return secoes


def renderizar(doc, corpo):
    linhas = corpo.splitlines()
    i = 0
    lista_num = 0
    primeira_secao = True

    while i < len(linhas):
        linha = linhas[i]
        crua = linha.strip()

        if not crua:
            lista_num = 0
            i += 1
            continue

        if crua == "---":
            fio(doc)
            i += 1
            continue

        if crua.startswith("## "):
            titulo = crua[3:].strip()
            if not primeira_secao:
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            primeira_secao = False
            m = re.match(r"^(\d+)\.\s*(.+)$", titulo)
            if m:
                faixa(doc, m.group(2).upper(), m.group(1))
            else:
                faixa(doc, titulo.upper())
            i += 1
            continue

        if crua.startswith("#### "):
            sub_subtitulo(doc, crua[5:].strip())
            i += 1
            continue

        if crua.startswith("### "):
            subtitulo(doc, crua[4:].strip())
            i += 1
            continue

        if crua.startswith("!> "):
            bloco = []
            while i < len(linhas) and linhas[i].strip().startswith("!> "):
                bloco.append(linhas[i].strip()[3:])
                i += 1
            caixa(doc, bloco, ROXO_CLARO, ROXO, rotulo="ATENÇÃO")
            continue

        if crua.startswith("> "):
            bloco = []
            while i < len(linhas) and linhas[i].strip().startswith("> "):
                bloco.append(linhas[i].strip()[2:])
                i += 1
            caixa(doc, bloco, DOURADO_CLARO, DOURADO)
            continue

        if linha_tabela(crua):
            bloco = []
            while i < len(linhas) and linha_tabela(linhas[i].strip()):
                bloco.append(linhas[i])
                i += 1
            quadro(doc, bloco)
            continue

        if re.match(r"^\s*-\s*\[\s*\]\s+", linha):
            recuo = len(linha) - len(linha.lstrip())
            texto = re.sub(r"^\s*-\s*\[\s*\]\s+", "", linha)
            if recuo >= 2:
                marcador(doc, texto, nivel=1)
            else:
                item_marcavel(doc, texto)
            i += 1
            continue

        if re.match(r"^\s*-\s+", linha):
            recuo = len(linha) - len(linha.lstrip())
            texto = re.sub(r"^\s*-\s+", "", linha)
            marcador(doc, texto, nivel=1 if recuo >= 2 else 0)
            i += 1
            continue

        m = re.match(r"^\s*(\d+)\.\s+(.*)$", linha)
        if m:
            lista_num += 1
            numerado(doc, m.group(1), m.group(2))
            i += 1
            continue

        paragrafo_comum(doc, crua)
        i += 1


# ---------------------------------------------------------------------------
# Principal
# ---------------------------------------------------------------------------


def gerar(caminho_md, saida=None, raiz=None):
    raiz = raiz or os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    meta, corpo = ler_fonte(caminho_md)
    logo = achar_logo(raiz)
    logo_escuro = achar_logo(raiz, CAMINHOS_LOGO_ESCURO) or logo

    doc = Document()
    estilo = doc.styles["Normal"]
    estilo.font.name = FONTE
    estilo.font.size = Pt(10.5)
    estilo.font.color.rgb = RGBColor.from_string(CINZA_TEXTO)

    secao = doc.sections[0]
    secao.top_margin = Cm(2.9)
    secao.bottom_margin = Cm(1.8)
    secao.left_margin = Cm(2.0)
    secao.right_margin = Cm(2.0)
    secao.header_distance = Cm(0.8)
    secao.footer_distance = Cm(0.9)

    montar_cabecalho_rodape(doc, meta, logo)
    montar_capa(doc, meta, logo_escuro)

    secoes = coletar_secoes(corpo)
    if secoes:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        montar_sumario(doc, secoes)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    renderizar(doc, corpo)

    fio(doc)
    p = doc.add_paragraph()
    espacamento(p, 6, 0, 1.2)
    r = p.add_run(
        f"Documento preparado por {MARCA} em {meta.get('data', '')}, "
        "com base no edital e nos anexos oficiais. "
        "Em caso de divergência, prevalece sempre o texto publicado pelo órgão."
    )
    r.italic = True
    r.font.size = Pt(8.5)
    r.font.name = FONTE
    r.font.color.rgb = RGBColor.from_string(CINZA_APOIO)

    if saida is None:
        saida = os.path.join(raiz, "Descricao Editais")
    if os.path.isdir(saida) or not saida.lower().endswith(".docx"):
        nome = meta.get("arquivo") or os.path.splitext(os.path.basename(caminho_md))[0]
        os.makedirs(saida, exist_ok=True)
        saida = os.path.join(saida, nome + ".docx")
    else:
        os.makedirs(os.path.dirname(saida), exist_ok=True)

    doc.save(saida)
    return saida, logo


def main():
    ap = argparse.ArgumentParser(description="Gera o Mapa do Edital em Word.")
    ap.add_argument("markdown", help="Arquivo markdown de origem")
    ap.add_argument("--saida", help="Pasta de destino ou caminho .docx", default=None)
    args = ap.parse_args()

    caminho, logo = gerar(args.markdown, args.saida)
    print("=== MAPA DO EDITAL ===")
    print("Arquivo: " + caminho)
    print("Logo: " + (logo if logo else "não encontrado, capa gerada com a marca em texto"))


if __name__ == "__main__":
    sys.stdout.reconfigure(encoding="utf-8")
    main()
